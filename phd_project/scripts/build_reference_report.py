#!/usr/bin/env python3
"""Build a Word report for selected references and HRC-related citing papers."""

from __future__ import annotations

import json
import os
import re
import textwrap
from datetime import datetime
from html import unescape
from typing import Any

import requests
from docx import Document

REFERENCE_URLS = [
    "https://www.semanticscholar.org/reader/7aca4e3ec6bc3b56169fb26670f1a98fc33f87ee",
    "https://dl.acm.org/doi/10.1145/3610977.3635006",
    "https://users.cs.utah.edu/~dsbrown/readings/learning_shared_autonomy.pdf",
    "https://www.mdpi.com/2218-6581/14/12/184",
    "https://journals.sagepub.com/doi/10.1177/02783649211050677",
]

OUTPUT_DOCX = "/home/g0amer/Desktop/thesis/outputs/reference_report_hrc.docx"
OUTPUT_JSON = "/home/g0amer/Desktop/thesis/outputs/reference_report_hrc.json"

HRC_KEYWORDS = [
    "human-robot",
    "human robot",
    "shared autonomy",
    "human-in-the-loop",
    "teleoperation",
    "collaborative robot",
    "cobot",
    "assistive robot",
    "human-aware",
    "hri",
    "hrc",
]

SESSION = requests.Session()
SESSION.headers.update(
    {
        "User-Agent": "ReferenceReportBot/1.0 (academic metadata collection)",
        "Accept": "application/json, text/plain, */*",
    }
)


def get_json(url: str, timeout: int = 25) -> dict[str, Any] | None:
    try:
        response = SESSION.get(url, timeout=timeout)
        if response.status_code != 200:
            return None
        return response.json()
    except Exception:
        return None


def get_text(url: str, timeout: int = 25) -> str | None:
    try:
        response = SESSION.get(url, timeout=timeout)
        if response.status_code != 200:
            return None
        return response.text
    except Exception:
        return None


def clean_text(value: str | None) -> str:
    if not value:
        return ""
    value = unescape(value)
    value = re.sub(r"<[^>]+>", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def extract_doi_from_url(url: str) -> str | None:
    doi_match = re.search(r"10\.\d{4,9}/[-._;()/:A-Za-z0-9]+", url)
    if doi_match:
        doi = doi_match.group(0).rstrip(".,;)>")
        return doi

    if "mdpi.com/2218-6581/14/12/184" in url:
        return "10.3390/robotics14120184"

    return None


def parse_crossref_date(message: dict[str, Any]) -> str:
    date_parts = (
        message.get("published-print", {}).get("date-parts")
        or message.get("published-online", {}).get("date-parts")
        or message.get("issued", {}).get("date-parts")
        or []
    )
    if not date_parts:
        return "Unknown"
    parts = date_parts[0]
    if len(parts) == 3:
        return f"{parts[0]:04d}-{parts[1]:02d}-{parts[2]:02d}"
    if len(parts) == 2:
        return f"{parts[0]:04d}-{parts[1]:02d}"
    if len(parts) == 1:
        return str(parts[0])
    return "Unknown"


def format_authors(author_list: list[dict[str, Any]] | None, max_authors: int = 8) -> str:
    if not author_list:
        return "Unknown"
    names = []
    for author in author_list[:max_authors]:
        given = author.get("given") or ""
        family = author.get("family") or ""
        name = (given + " " + family).strip()
        if not name:
            name = author.get("name") or ""
        if name:
            names.append(name)
    if len(author_list) > max_authors:
        names.append("et al.")
    return ", ".join(names) if names else "Unknown"


def summarize_scientific_details(abstract: str, title: str) -> str:
    text = clean_text(abstract)
    if not text:
        return "No abstract retrieved; scientific details could not be extracted automatically."

    lowered = text.lower()
    signals = []

    if any(k in lowered for k in ["dataset", "corpus", "benchmark"]):
        signals.append("The work includes or introduces a dataset/benchmark.")
    if any(k in lowered for k in ["framework", "pipeline", "architecture", "system"]):
        signals.append("The work proposes a framework or system architecture.")
    if any(k in lowered for k in ["model", "learning", "reinforcement learning", "classification", "prediction"]):
        signals.append("The work develops or evaluates machine learning methods.")
    if any(k in lowered for k in ["experiment", "user study", "participants", "subjects"]):
        signals.append("The work reports experimental evaluation with human participants or task trials.")
    if any(k in lowered for k in ["eeg", "ecg", "eda", "emg", "physiological", "biosignal"]):
        signals.append("The work uses physiological/biosignal modalities.")

    first_sentences = re.split(r"(?<=[.!?])\s+", text)
    synopsis = " ".join(first_sentences[:2]).strip()
    synopsis = textwrap.shorten(synopsis, width=450, placeholder="...")

    if not signals:
        signals.append("The abstract indicates a methodological contribution relevant to HRC/HRI.")

    return synopsis + " " + " ".join(signals)


def get_crossref_metadata(doi: str) -> dict[str, Any] | None:
    url = f"https://api.crossref.org/works/{doi}"
    data = get_json(url)
    if not data or "message" not in data:
        return None
    msg = data["message"]
    title = ""
    if isinstance(msg.get("title"), list) and msg["title"]:
        title = msg["title"][0]
    abstract = clean_text(msg.get("abstract", ""))
    return {
        "title": title or "Unknown title",
        "doi": doi,
        "release_date": parse_crossref_date(msg),
        "venue": (msg.get("container-title") or ["Unknown venue"])[0],
        "authors": format_authors(msg.get("author")),
        "abstract": abstract,
        "source": "Crossref",
    }


def find_openalex_work(doi: str | None = None, title_query: str | None = None) -> dict[str, Any] | None:
    if doi:
        encoded = requests.utils.quote(f"https://doi.org/{doi}", safe="")
        url = f"https://api.openalex.org/works/{encoded}"
        data = get_json(url)
        if data and data.get("id"):
            return data

    if title_query:
        query = requests.utils.quote(title_query)
        url = f"https://api.openalex.org/works?search={query}&per-page=5"
        data = get_json(url)
        if data and data.get("results"):
            return data["results"][0]

    return None


def get_semantic_scholar_paper(paper_id_or_query: str, is_query: bool = False) -> dict[str, Any] | None:
    if not is_query:
        fields = "title,year,abstract,externalIds,url,venue,authors,publicationDate,citationCount"
        url = f"https://api.semanticscholar.org/graph/v1/paper/{paper_id_or_query}?fields={fields}"
        return get_json(url)

    fields = "title,year,abstract,externalIds,url,venue,authors,publicationDate,citationCount"
    query = requests.utils.quote(paper_id_or_query)
    url = f"https://api.semanticscholar.org/graph/v1/paper/search?query={query}&limit=5&fields={fields}"
    data = get_json(url)
    if not data or not data.get("data"):
        return None

    preferred = None
    for candidate in data["data"]:
        t = (candidate.get("title") or "").lower()
        if "shared autonomy" in t:
            preferred = candidate
            break
    return preferred or data["data"][0]


def normalize_seed_from_semantic(entry: dict[str, Any], source_url: str) -> dict[str, Any]:
    external_ids = entry.get("externalIds") or {}
    doi = external_ids.get("DOI")
    return {
        "title": entry.get("title") or "Unknown title",
        "doi": doi,
        "release_date": entry.get("publicationDate") or str(entry.get("year") or "Unknown"),
        "venue": entry.get("venue") or "Unknown venue",
        "authors": format_authors(entry.get("authors")),
        "abstract": clean_text(entry.get("abstract", "")),
        "semantic_paper_id": entry.get("paperId"),
        "source_url": source_url,
        "source": "Semantic Scholar",
    }


def hrc_relevant(text: str) -> bool:
    lower = (text or "").lower()
    return any(keyword in lower for keyword in HRC_KEYWORDS)


def parse_openalex_citation_item(item: dict[str, Any]) -> dict[str, Any]:
    authors = ", ".join(
        [
            (a.get("author", {}).get("display_name") or "")
            for a in item.get("authorships", [])[:8]
            if a.get("author", {}).get("display_name")
        ]
    ) or "Unknown"

    abstract = ""
    inv = item.get("abstract_inverted_index")
    if isinstance(inv, dict) and inv:
        token_positions: list[tuple[int, str]] = []
        for token, positions in inv.items():
            for pos in positions:
                token_positions.append((pos, token))
        token_positions.sort(key=lambda x: x[0])
        abstract = " ".join(token for _, token in token_positions)

    return {
        "title": item.get("display_name") or "Unknown title",
        "release_date": item.get("publication_date") or str(item.get("publication_year") or "Unknown"),
        "venue": (item.get("primary_location") or {}).get("source", {}).get("display_name") or "Unknown venue",
        "doi": (item.get("ids") or {}).get("doi", "").replace("https://doi.org/", "") or None,
        "authors": authors,
        "abstract": clean_text(abstract),
        "source": "OpenAlex",
    }


def get_hrc_citing_papers(seed: dict[str, Any], max_items: int = 4) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []

    # 1) OpenAlex cited-by route if possible.
    openalex_work = None
    if seed.get("doi"):
        openalex_work = find_openalex_work(doi=seed["doi"])
    if not openalex_work and seed.get("title"):
        openalex_work = find_openalex_work(title_query=seed["title"])

    if openalex_work and openalex_work.get("cited_by_api_url"):
        cited_api = openalex_work["cited_by_api_url"] + "?per-page=50"
        data = get_json(cited_api)
        for item in (data or {}).get("results", []):
            parsed = parse_openalex_citation_item(item)
            hay = (parsed.get("title", "") + " " + parsed.get("abstract", "")).lower()
            if hrc_relevant(hay):
                candidates.append(parsed)

    # 2) Semantic Scholar fallback.
    sem_paper = None
    if seed.get("semantic_paper_id"):
        sem_paper = seed.get("semantic_paper_id")
    elif seed.get("doi"):
        sem_paper = f"DOI:{seed['doi']}"

    if sem_paper and len(candidates) < max_items:
        fields = (
            "citingPaper.title,citingPaper.year,citingPaper.abstract,citingPaper.externalIds,"
            "citingPaper.url,citingPaper.venue,citingPaper.authors,citingPaper.publicationDate"
        )
        url = (
            f"https://api.semanticscholar.org/graph/v1/paper/{requests.utils.quote(sem_paper, safe='')}/"
            f"citations?fields={fields}&limit=100"
        )
        data = get_json(url)
        for row in (data or {}).get("data", []):
            cp = row.get("citingPaper") or {}
            title = cp.get("title") or "Unknown title"
            abstract = clean_text(cp.get("abstract", ""))
            hay = (title + " " + abstract).lower()
            if not hrc_relevant(hay):
                continue
            ext = cp.get("externalIds") or {}
            citem = {
                "title": title,
                "release_date": cp.get("publicationDate") or str(cp.get("year") or "Unknown"),
                "venue": cp.get("venue") or "Unknown venue",
                "doi": ext.get("DOI"),
                "authors": format_authors(cp.get("authors")),
                "abstract": abstract,
                "source": "Semantic Scholar",
            }
            candidates.append(citem)

    deduped: list[dict[str, Any]] = []
    seen = set()
    for item in candidates:
        key = (item.get("title") or "").strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        deduped.append(item)

    return deduped[:max_items]


def build_seed_record(source_url: str) -> dict[str, Any]:
    doi = extract_doi_from_url(source_url)

    # Hard-coded resolved entries (from manual arXiv + CrossRef lookups).
    if "semanticscholar.org/reader/7aca4e3ec6bc3b56169fb26670f1a98fc33f87ee" in source_url:
        return {
            "title": "Kaiwu: A Multimodal Manipulation Dataset and Framework for Robot Learning and Human-Robot Interaction",
            "doi": "10.48550/arXiv.2503.05231",
            "release_date": "2025-03-07",
            "venue": "arXiv (Submitted to IEEE Robotics and Automation Letters)",
            "authors": "Shuo Jiang, Haonan Li, Ruochen Ren, Yanmin Zhou, Zhipeng Wang, Bin He",
            "abstract": "This paper presents the Kaiwu multimodal dataset to address missing real-world synchronized multimodal data in sophisticated assembling scenarios. The dataset integrates human, environment and robot data from 20 subjects and 30 interaction objects, yielding 11,664 integrated action instances. For each demonstration, hand motions, operation pressures, assembly sounds, multi-view videos, high-precision motion capture, eye gaze with first-person videos, and electromyography signals are recorded. The dataset aims to facilitate robot learning, dexterous manipulation, human intention investigation, and human-robot collaboration research.",
            "source_url": source_url,
            "source": "arXiv",
        }

    if "learning_shared_autonomy.pdf" in source_url:
        return {
            "title": "Learning to Share Autonomy from Repeated Human-Robot Interaction",
            "doi": "10.1109/iros51168.2021.9636748",
            "release_date": "2021",
            "venue": "2021 IEEE/RSJ International Conference on Intelligent Robots and Systems (IROS)",
            "authors": "Ananth Jonnavittula, Shaunak A. Mehta, Dylan P. Losey",
            "abstract": "This paper proposes an approach to shared autonomy that learns assistance from scratch rather than relying on prior knowledge. The key insight is that operators repeat important tasks daily (e.g., opening the fridge, making coffee). By learning from repeated interactions, the robot can improve its understanding of user intent over time and provide increasingly effective assistance in shared autonomy scenarios.",
            "source_url": source_url,
            "source": "IEEE",
        }

    if "semanticscholar.org/reader/" in source_url:
        paper_id = source_url.rstrip("/").split("/")[-1]
        sem = get_semantic_scholar_paper(paper_id)
        if sem:
            rec = normalize_seed_from_semantic(sem, source_url)
            if not rec.get("doi"):
                rec["doi"] = extract_doi_from_url(sem.get("url", "") or "")
            return rec

    if doi:
        cross = get_crossref_metadata(doi)
        if cross:
            cross["source_url"] = source_url
            return cross

    return {
        "title": "Unresolved reference",
        "doi": doi,
        "release_date": "Unknown",
        "venue": "Unknown",
        "authors": "Unknown",
        "abstract": "Could not resolve metadata from available APIs.",
        "source_url": source_url,
        "source": "Unresolved",
    }


def write_docx(records: list[dict[str, Any]]) -> None:
    doc = Document()
    doc.add_heading("Reference Analysis for Human-Robot Collaboration", 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(
        "This report includes: release date, scientific details (data/framework), and HRC-related papers that cite each seed reference."
    )

    for idx, rec in enumerate(records, start=1):
        doc.add_heading(f"{idx}. {rec.get('title', 'Unknown title')}", level=1)
        doc.add_paragraph(f"Source URL: {rec.get('source_url', 'N/A')}")
        doc.add_paragraph(f"Release date: {rec.get('release_date', 'Unknown')}")
        doc.add_paragraph(f"Venue: {rec.get('venue', 'Unknown')}")
        doc.add_paragraph(f"DOI: {rec.get('doi') or 'N/A'}")
        doc.add_paragraph(f"Authors: {rec.get('authors', 'Unknown')}")

        science = summarize_scientific_details(rec.get("abstract", ""), rec.get("title", ""))
        doc.add_paragraph("Scientific details:")
        doc.add_paragraph(science)

        citing = rec.get("hrc_citing_papers", [])
        doc.add_heading("HRC-related papers citing this work", level=2)
        if not citing:
            doc.add_paragraph("No clearly HRC-related citing papers were automatically identified from accessible metadata sources.")
            continue

        for cidx, cp in enumerate(citing, start=1):
            doc.add_heading(f"{idx}.{cidx} {cp.get('title', 'Unknown title')}", level=3)
            doc.add_paragraph(f"Release date: {cp.get('release_date', 'Unknown')}")
            doc.add_paragraph(f"Venue: {cp.get('venue', 'Unknown')}")
            doc.add_paragraph(f"DOI: {cp.get('doi') or 'N/A'}")
            doc.add_paragraph(f"Authors: {cp.get('authors', 'Unknown')}")
            details = summarize_scientific_details(cp.get("abstract", ""), cp.get("title", ""))
            doc.add_paragraph("Scientific details:")
            doc.add_paragraph(details)

    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
    doc.save(OUTPUT_DOCX)


def main() -> None:
    records: list[dict[str, Any]] = []

    for ref in REFERENCE_URLS:
        seed = build_seed_record(ref)
        seed["hrc_citing_papers"] = get_hrc_citing_papers(seed)
        records.append(seed)

    os.makedirs(os.path.dirname(OUTPUT_JSON), exist_ok=True)
    with open(OUTPUT_JSON, "w", encoding="utf-8") as f:
        json.dump(records, f, indent=2, ensure_ascii=False)

    write_docx(records)
    print(f"Saved: {OUTPUT_DOCX}")
    print(f"Saved: {OUTPUT_JSON}")


if __name__ == "__main__":
    main()
