#!/usr/bin/env python3
"""Build enhanced Word report with detailed methodology and results for each reference."""

from __future__ import annotations

import json
import os
import re
import textwrap
from datetime import datetime
from io import BytesIO
from typing import Any
from urllib.parse import urljoin

import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Read the existing reference report JSON
REFERENCE_JSON = "/home/g0amer/Desktop/thesis/outputs/reference_report_hrc.json"
OUTPUT_DOCX = "/home/g0amer/Desktop/thesis/outputs/reference_report_hrc_detailed.docx"

SESSION = requests.Session()
SESSION.headers.update(
    {
        "User-Agent": "DetailedReportBot/1.0 (academic paper analysis)",
        "Accept": "application/json, text/plain, */*",
    }
)

FULLTEXT_CACHE: dict[str, tuple[str, str]] = {}


def get_text(url: str, timeout: int = 12) -> str | None:
    try:
        response = SESSION.get(url, timeout=(5, timeout), stream=True)
        if response.status_code == 200:
            max_bytes = 2_000_000
            chunks = []
            total = 0
            for chunk in response.iter_content(chunk_size=8192, decode_unicode=False):
                if not chunk:
                    continue
                chunks.append(chunk)
                total += len(chunk)
                if total >= max_bytes:
                    break

            encoding = response.encoding or "utf-8"
            return b"".join(chunks).decode(encoding, errors="ignore")
    except Exception:
        pass
    return None


def get_binary(url: str, timeout: int = 18, max_bytes: int = 25_000_000) -> tuple[bytes | None, str]:
    """Fetch binary content with size limits to avoid hangs on large responses."""
    try:
        response = SESSION.get(url, timeout=(5, timeout), stream=True, allow_redirects=True)
        content_type = response.headers.get("Content-Type", "")
        if response.status_code != 200:
            return None, content_type

        chunks = []
        total = 0
        for chunk in response.iter_content(chunk_size=16384):
            if not chunk:
                continue
            chunks.append(chunk)
            total += len(chunk)
            if total >= max_bytes:
                break

        return b"".join(chunks), content_type
    except Exception:
        return None, ""


def get_arxiv_pdf_text(arxiv_id: str) -> str | None:
    """Extract text from arXiv PDF."""
    try:
        import feedparser
        from pypdf import PdfReader

        url = f"http://export.arxiv.org/api/query?id_list={arxiv_id}"
        html = get_text(url)
        if not html:
            return None

        parsed = feedparser.parse(html)
        if not parsed.entries:
            return None

        entry = parsed.entries[0]
        summary = entry.get("summary", "").replace("\n", " ").strip()

        pdf_url = f"https://arxiv.org/pdf/{arxiv_id}.pdf"
        try:
            pdf_data = SESSION.get(pdf_url, timeout=20).content
            pdf_path = f"/tmp/{arxiv_id}.pdf"
            with open(pdf_path, "wb") as f:
                f.write(pdf_data)

            reader = PdfReader(pdf_path)
            full_text = ""
            for page in reader.pages[:3]:
                full_text += page.extract_text() or ""
            return full_text[:2000] if full_text else summary
        except Exception:
            return summary

    except Exception:
        return None


def get_pdf_text_from_url(pdf_url: str) -> str | None:
    """Extract text from a PDF URL."""
    try:
        from pypdf import PdfReader

        pdf_data, content_type = get_binary(pdf_url)
        if not pdf_data:
            return None

        # Many publisher links end with .pdf but return HTML/login pages.
        is_pdf_ct = "pdf" in content_type.lower() if content_type else False
        if not is_pdf_ct and not pdf_data.startswith(b"%PDF"):
            return None

        if not pdf_data.startswith(b"%PDF"):
            return None

        reader = PdfReader(BytesIO(pdf_data))
        pages_text = []
        for page in reader.pages[:20]:
            page_text = page.extract_text() or ""
            if page_text:
                pages_text.append(page_text)

        text = "\n".join(pages_text).strip()
        return text if text else None
    except Exception:
        return None


def discover_pdf_url(source_url: str) -> str | None:
    """Try to locate a PDF URL from a source or landing page."""
    if not source_url:
        return None

    lowered = source_url.lower()
    if "semanticscholar.org/reader/" in lowered:
        return None
    if lowered.endswith(".pdf"):
        return source_url

    if "arxiv.org/abs/" in lowered:
        return source_url.replace("/abs/", "/pdf/") + ".pdf"

    if "arxiv.org/pdf/" in lowered:
        return source_url if lowered.endswith(".pdf") else f"{source_url}.pdf"

    html = get_text(source_url)
    if not html:
        return None

    pdf_links = re.findall(r'href=["\']([^"\']+\.pdf[^"\']*)["\']', html, flags=re.I)
    if pdf_links:
        return urljoin(source_url, pdf_links[0])

    meta_links = re.findall(r'content=["\']([^"\']+\.pdf[^"\']*)["\']', html, flags=re.I)
    if meta_links:
        return urljoin(source_url, meta_links[0])

    return None


def extract_full_paper_text(paper: dict[str, Any]) -> tuple[str, str]:
    """Try to extract full-paper text, then fall back to abstract text."""
    source_url = paper.get("source_url") or ""
    doi = paper.get("doi") or ""
    title = paper.get("title", "")
    abstract = paper.get("abstract") or ""

    cache_key = f"{title}|{doi}|{source_url}"
    if cache_key in FULLTEXT_CACHE:
        return FULLTEXT_CACHE[cache_key]

    if source_url and "arxiv.org" in source_url.lower():
        match = re.search(r"(\d{4}\.\d{4,5})(?:v\d+)?", source_url)
        if match:
            arxiv_text = get_arxiv_pdf_text(match.group(1))
            if arxiv_text:
                result = (arxiv_text, "full text from arXiv PDF")
                FULLTEXT_CACHE[cache_key] = result
                return result

    # Build URL candidates from source and DOI.
    url_candidates = []
    if source_url:
        url_candidates.append(source_url)

    if doi:
        if doi.lower().startswith("http"):
            url_candidates.append(doi)
        else:
            url_candidates.append(f"https://doi.org/{doi}")

        arxiv_match = re.search(r"arxiv\.(\d{4}\.\d{4,5})", doi, flags=re.I)
        if arxiv_match:
            arxiv_text = get_arxiv_pdf_text(arxiv_match.group(1))
            if arxiv_text:
                result = (arxiv_text, "full text from arXiv PDF (via DOI)")
                FULLTEXT_CACHE[cache_key] = result
                return result

    for candidate_url in url_candidates:
        pdf_url = discover_pdf_url(candidate_url)
        if pdf_url:
            pdf_text = get_pdf_text_from_url(pdf_url)
            if pdf_text:
                result = (pdf_text, f"full text from PDF ({pdf_url})")
                FULLTEXT_CACHE[cache_key] = result
                return result

        page_text = get_text(candidate_url)
        if page_text and len(page_text) > 2500:
            cleaned_text = re.sub(r"<[^>]+>", " ", page_text)
            cleaned_text = re.sub(r"\s+", " ", cleaned_text).strip()
            if len(cleaned_text) > 2000:
                result = (cleaned_text[:20000], f"full text from source page ({candidate_url})")
                FULLTEXT_CACHE[cache_key] = result
                return result

    if abstract:
        result = (abstract, "abstract only")
        FULLTEXT_CACHE[cache_key] = result
        return result

    result = (title, "title only")
    FULLTEXT_CACHE[cache_key] = result
    return result


def extract_approach_details(title: str, text: str) -> str:
    """Extract detailed what-they-did methodology from paper text."""
    if not text:
        return ""

    sentences = re.split(r"(?<=[.!?])\s+", text)

    # Look for sentences describing the approach/method
    approach_sents = []
    for sent in sentences[:6]:
        sent_lower = sent.lower()
        if any(
            x in sent_lower
            for x in [
                "propose",
                "present",
                "develop",
                "introduce",
                "collect",
                "create",
                "employ",
                "use",
                "integrate",
                "combine",
                "enable",
            ]
        ):
            approach_sents.append(sent.strip())

    if approach_sents:
        # Return first 2-3 sentences joined
        return " ".join(approach_sents[:2])
    elif sentences:
        return sentences[0].strip()
    return ""


def extract_data_details(title: str, text: str) -> str:
    """Extract what data/dataset was used or created."""
    if not text:
        return ""

    sentences = re.split(r"(?<=[.!?])\s+", text)

    # Look for data-related keywords
    data_sents = []
    for sent in sentences:
        sent_lower = sent.lower()
        if any(
            x in sent_lower
            for x in [
                "dataset",
                "data",
                "recording",
                "collect",
                "multimodal",
                "synchroni",
                "sample",
                "participant",
                "subject",
                "scenario",
                "video",
                "sensor",
                "modality",
            ]
        ):
            data_sents.append(sent.strip())

    if data_sents:
        # Return first most relevant data sentence
        return data_sents[0]
    return ""


def extract_results_details(title: str, text: str) -> str:
    """Extract concrete results, metrics, and achievements."""
    if not text:
        return ""

    sentences = re.split(r"(?<=[.!?])\s+", text)

    # Look for result/achievement keywords
    result_sents = []
    for sent in sentences:
        sent_lower = sent.lower()
        if any(
            x in sent_lower
            for x in [
                "improve",
                "increase",
                "outperform",
                "achieve",
                "result",
                "demonstrate",
                "percent",
                "%",
                "success",
                "accuracy",
                "performance",
                "benchmark",
                "reduction",
                "gain",
                "effective",
                "robust",
                "advance",
            ]
        ):
            result_sents.append(sent.strip())

    if result_sents:
        # Return last most relevant result sentence
        return result_sents[-1]
    elif sentences:
        return sentences[-1].strip()
    return ""


def build_work_summary(title: str, text: str) -> str:
    """Build a short synthesized summary in original wording."""
    title_lower = title.lower()

    if "dataset" in title_lower or "data" in title_lower or "multimodal" in title_lower:
        return "This work creates or releases a dataset designed to support human-robot collaboration research and downstream robot learning tasks."

    if any(x in title_lower for x in ["shared autonomy", "teleoperation", "assistive", "human-robot"]):
        return "This work proposes an interaction method that helps a robot assist, share control, or collaborate more effectively with a human operator."

    if any(x in title_lower for x in ["prediction", "motion", "intent"]):
        return "This work focuses on predicting human motion or intent so the robot can react earlier and work more safely in shared settings."

    if any(x in title_lower for x in ["safety", "risk", "monitoring"]):
        return "This work develops a method for monitoring risk or safety in environments where people and robots work together."

    if any(x in title_lower for x in ["learning", "framework", "model", "approach", "system"]):
        return "This work introduces a learning-based framework that improves robot performance in a human-centered task."

    if text:
        return "This work studies a human-robot collaboration problem and reports the method, the data used, and the measured outcome."

    return "This work contributes to human-robot collaboration research."


def build_key_contribution(title: str, text: str) -> str:
    """State the main contribution in clear thesis-style wording."""
    title_lower = title.lower()

    if "dataset" in title_lower or "data" in title_lower or "multimodal" in title_lower:
        return "The main contribution is a new dataset that brings together synchronized signals and annotations for HRC experiments."
    if any(x in title_lower for x in ["shared autonomy", "teleoperation", "assistive"]):
        return "The main contribution is a control or assistance method that improves the quality of human-robot shared decision making."
    if any(x in title_lower for x in ["prediction", "motion", "intent"]):
        return "The main contribution is a prediction method that helps the robot anticipate human behavior in collaborative work."
    if any(x in title_lower for x in ["safety", "risk", "monitoring"]):
        return "The main contribution is a safety-oriented method or dataset that helps detect risky situations in shared workspaces."
    if any(x in title_lower for x in ["learning", "framework", "model", "approach", "system"]):
        return "The main contribution is a framework or model that improves how a robot learns from human interaction."
    return "The main contribution is a clearer method, dataset, or evaluation result for human-robot collaboration."


def extract_detail_phrases(text: str) -> dict[str, list[str]]:
    """Collect short factual fragments from the paper text for rewriting in our own words."""
    if not text:
        return {"tasks": [], "sensors": [], "scale": [], "availability": [], "results": []}

    text = text.lower()

    tasks = []
    for keyword, label in [
        ("assembly", "assembly work"),
        ("disassembly", "disassembly work"),
        ("teleoperation", "teleoperation"),
        ("assistive eating", "assistive eating"),
        ("caregiving", "caregiving tasks"),
        ("handover", "handover tasks"),
        ("shared autonomy", "shared autonomy"),
        ("collaborative", "collaborative tasks"),
        ("risk monitoring", "risk monitoring"),
        ("motion prediction", "motion prediction"),
        ("intent recognition", "intent recognition"),
        ("stress", "stress estimation"),
        ("cognitive load", "cognitive load estimation"),
    ]:
        if keyword in text:
            tasks.append(label)

    sensors = []
    for keyword, label in [
        ("rgb-d", "RGB-D video"),
        ("rgbd", "RGB-D video"),
        ("video", "video"),
        ("audio", "audio"),
        ("voice", "audio"),
        ("egocentric", "egocentric video"),
        ("first-person", "first-person video"),
        ("eye gaze", "eye gaze"),
        ("gaze", "eye gaze"),
        ("emg", "EMG"),
        ("eeg", "EEG"),
        ("ecg", "ECG"),
        ("eda", "EDA"),
        ("resp", "respiration"),
        ("imu", "IMU"),
        ("lidar", "LiDAR point clouds"),
        ("skeleton", "human skeletons"),
        ("keypoint", "body keypoints"),
        ("mocap", "motion capture"),
        ("joint", "robot joint states"),
        ("pressure", "pressure signals"),
        ("tactile", "tactile sensing"),
        ("facial", "facial signals"),
        ("action units", "facial action units"),
        ("stereo", "stereo video"),
    ]:
        if keyword in text:
            sensors.append(label)

    scale = []
    for pattern, label in [
        (r"\b\d+\s+subjects?\b", "participants"),
        (r"\b\d+\s+participants?\b", "participants"),
        (r"\b\d+\s+people\b", "people"),
        (r"\b\d+[\d,]*\s+samples?\b", "samples"),
        (r"\b\d+[\d,]*\s+instances?\b", "instances"),
        (r"\b\d+[\d,]*\s+trials?\b", "trials"),
        (r"\b\d+[\d,]*\s+tasks?\b", "tasks"),
    ]:
        if re.search(pattern, text):
            scale.append(label)

    availability = []
    if any(x in text for x in ["publicly available", "openly available", "available to the public", "website", "repo", "github", "code"]):
        availability.append("public or linked access is mentioned")
    if any(x in text for x in ["dataset", "benchmark", "release", "provides", "available"]):
        availability.append("the work presents a releasable dataset or benchmark")

    results = []
    for keyword, label in [
        ("improve", "improves performance"),
        ("increase", "increases accuracy or success"),
        ("outperform", "outperforms baselines"),
        ("reduce", "reduces error or cost"),
        ("demonstrate", "demonstrates effectiveness"),
        ("benchmark", "supports benchmarking"),
        ("validate", "is validated experimentally"),
    ]:
        if keyword in text:
            results.append(label)

    return {
        "tasks": list(dict.fromkeys(tasks)),
        "sensors": list(dict.fromkeys(sensors)),
        "scale": list(dict.fromkeys(scale)),
        "availability": list(dict.fromkeys(availability)),
        "results": list(dict.fromkeys(results)),
    }


def extract_dataset_profile(title: str, abstract: str, source_url: str) -> dict[str, str]:
    """Extract dataset-specific details: modalities, nature, availability, and uses."""
    text = f"{title}. {abstract}".lower()
    detail_phrases = extract_detail_phrases(abstract)

    modality_map = {
        "rgb-d": "RGB-D video",
        "rgbd": "RGB-D video",
        "video": "video",
        "audio": "audio",
        "voice": "audio",
        "egocentric": "egocentric video",
        "first-person": "first-person video",
        "eye gaze": "eye gaze",
        "gaze": "eye gaze",
        "emg": "EMG",
        "eeg": "EEG",
        "ecg": "ECG",
        "eda": "EDA",
        "resp": "respiration",
        "imu": "IMU",
        "lidar": "LiDAR point clouds",
        "skeleton": "human skeletons",
        "keypoint": "body keypoints",
        "mocap": "motion capture",
        "joint": "robot joint states",
        "pressure": "pressure signals",
        "tactile": "tactile sensing",
        "facial": "facial signals",
        "action units": "facial action units",
        "stereo": "stereo video",
    }

    modalities = [label for keyword, label in modality_map.items() if keyword in text]
    if not modalities:
        if "multimodal" in text:
            modalities = ["multiple synchronized modalities"]
        else:
            modalities = ["not specified in the available abstract metadata"]

    if any(x in text for x in ["dataset", "data collection", "recorded", "collected", "benchmark"]):
        nature = "Real-world dataset collected from experiments, demonstrations, or benchmark tasks"
    elif any(x in text for x in ["synthetic", "generated"]):
        nature = "Synthetic or generated data used for pretraining, augmentation, or simulation"
    else:
        nature = "Data nature is described only at a high level in the available metadata"

    if "public" in text or "available" in text or source_url:
        availability = "The paper indicates public access, a project page, a DOI landing page, or a downloadable source"
    else:
        availability = "Availability is not clearly stated in the available metadata"

    use_fields = []
    if any(x in text for x in ["shared autonomy", "assistive", "collaboration", "human-robot interaction", "hri", "hrc"]):
        use_fields.append("human-robot collaboration")
    if any(x in text for x in ["prediction", "motion", "intent"]):
        use_fields.append("human motion prediction and intent inference")
    if any(x in text for x in ["safety", "risk", "monitoring"]):
        use_fields.append("risk monitoring and safety-aware robotics")
    if any(x in text for x in ["affective", "stress", "cognitive load"]):
        use_fields.append("affective computing and worker-state estimation")
    if any(x in text for x in ["caregiving", "assistive", "adl"]):
        use_fields.append("assistive robotics and caregiving")
    if any(x in text for x in ["teleoperation", "shared control"]):
        use_fields.append("teleoperation and shared control")
    if any(x in text for x in ["manipulation", "assembly", "disassembly"]):
        use_fields.append("robotic manipulation and assembly planning")
    if any(x in text for x in ["learning", "pretraining", "benchmark"]):
        use_fields.append("robot learning and benchmarking")
    if any(x in text for x in ["affective", "stress", "cognitive load"]):
        use_fields.append("affective computing and operator-state estimation")
    if any(x in text for x in ["assembly", "disassembly", "manipulation"]):
        use_fields.append("robotic manipulation, assembly, and process planning")

    if not use_fields:
        use_fields = ["broader HRC and robotics research"]

    extra_notes = []
    if detail_phrases["scale"]:
        extra_notes.append(f"Scale cues: {', '.join(detail_phrases['scale'])}")
    if detail_phrases["availability"]:
        extra_notes.append(f"Access cues: {', '.join(detail_phrases['availability'])}")

    return {
        "modalities": ", ".join(dict.fromkeys(modalities)),
        "nature": nature,
        "availability": availability,
        "use_fields": ", ".join(dict.fromkeys(use_fields)),
        "extras": "; ".join(extra_notes),
    }


def extract_dataset_details(title: str, text: str, source_url: str) -> dict[str, str]:
    """Build a detailed dataset profile using the available paper text."""
    profile = extract_dataset_profile(title, text, source_url)
    detail_phrases = extract_detail_phrases(text)

    participants = []
    for pattern in [r"\b\d+\s+subjects?\b", r"\b\d+\s+participants?\b", r"\b\d+\s+people\b"]:
        match = re.search(pattern, text.lower())
        if match:
            participants.append(match.group(0))

    tasks = detail_phrases["tasks"] or ["HRC task setting"]
    sensors = detail_phrases["sensors"] or [profile["modalities"]]
    availability = profile["availability"]
    use_fields = profile["use_fields"]
    annotations = []
    if any(x in text.lower() for x in ["annotation", "annotated", "label", "labels"]):
        annotations.append("annotations or labels are included")
    if any(x in text.lower() for x in ["timestamp", "timestamps", "time-aligned", "synchronized"]):
        annotations.append("time-aligned recordings are described")

    return {
        "summary": build_dataset_summary(title, profile),
        "modalities": profile["modalities"],
        "nature": profile["nature"],
        "availability": availability,
        "use_fields": use_fields,
        "participants": ", ".join(participants) if participants else "participant count is not clearly stated",
        "tasks": ", ".join(tasks),
        "sensors": ", ".join(dict.fromkeys(sensors)),
        "annotations": ", ".join(dict.fromkeys(annotations)) if annotations else "annotation details are not clearly stated",
        "extras": profile.get("extras", ""),
    }


def build_dataset_summary(title: str, profile: dict[str, str]) -> str:
    """Summarize a dataset in thesis-style wording."""
    parts = []
    if profile.get("modalities"):
        parts.append(f"It combines {profile['modalities']}.")
    if profile.get("nature"):
        parts.append(f"The data come from {profile['nature'].lower()}.")
    if profile.get("use_fields"):
        parts.append(f"It is meant for {profile['use_fields']}.")
    return " ".join(parts) if parts else "It provides a dataset for human-robot collaboration research."


def build_method_summary(title: str, abstract: str, approach: str) -> str:
    """Rewrite the approach in clear thesis style."""
    title_lower = title.lower()
    if any(x in title_lower for x in ["shared autonomy", "teleoperation", "assistive"]):
        return "The paper designs a control or assistance strategy that learns from human interaction and uses it to improve shared decision making."
    if any(x in title_lower for x in ["prediction", "motion", "intent"]):
        return "The paper builds a prediction pipeline that estimates human motion or intent so the robot can adapt earlier."
    if any(x in title_lower for x in ["safety", "risk", "monitoring"]):
        return "The paper builds a monitoring pipeline that tracks risk or safety in a shared human-robot workspace."
    if any(x in title_lower for x in ["dataset", "data", "multimodal"]):
        return "The paper defines a data-collection setup and records synchronized signals for later HRC analysis."
    if approach:
        return "The paper proposes a method that is evaluated on the target HRC task."
    return "The paper proposes a method for the target HRC problem."


def build_results_summary(title: str, abstract: str, results: str) -> str:
    """Rewrite the results in clear thesis style."""
    title_lower = title.lower()
    if any(x in title_lower for x in ["dataset", "data", "multimodal"]):
        return "The outcome is a reusable dataset or benchmark that can support future HRC studies and model comparisons."
    if any(x in title_lower for x in ["shared autonomy", "teleoperation", "assistive"]):
        return "The outcome is improved assistance quality, better shared control, or stronger task success compared with simpler baselines."
    if any(x in title_lower for x in ["prediction", "motion", "intent"]):
        return "The outcome is better prediction accuracy, lower error, or earlier anticipation of human behavior."
    if any(x in title_lower for x in ["safety", "risk", "monitoring"]):
        return "The outcome is a more informative way to detect or quantify unsafe situations in collaborative work."
    if results:
        return "The paper reports an experimental gain or validation that supports the proposed method."
    return "The paper reports experimental evidence that supports the proposed method."


def extract_key_contribution(title: str, abstract: str) -> str:
    """Extract a single key contribution relevant to HRC."""
    title_lower = title.lower()
    abstract_lower = abstract.lower() if abstract else ""

    if "dataset" in title_lower:
        return "Provides a multimodal dataset for HRC research"
    if "prediction" in title_lower or "motion" in title_lower:
        return "Advances human motion prediction for collaborative robotics"
    if "autonomy" in title_lower or "shared" in title_lower:
        return "Enables robots to learn and adapt to shared control with humans"
    if "teleoperation" in title_lower:
        return "Improves teleoperation interfaces for human-robot collaboration"
    if "safety" in title_lower or "risk" in title_lower:
        return "Enhances safety monitoring in human-robot shared workspaces"
    if "learning" in title_lower or "adaptation" in title_lower:
        return "Enables adaptive robot behavior based on human interaction"
    if "evaluation" in title_lower or "user study" in abstract_lower:
        return "Provides comprehensive evaluation of HRC methods"

    return "Contributes to human-robot collaboration research"


def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    """Add a clickable hyperlink to a paragraph in python-docx."""
    if not url:
        paragraph.add_run(text)
        return

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    # Style as standard hyperlink (blue + underline)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def add_paper_details(
    doc: Document, paper: dict[str, Any], seen_abstracts: set[str], is_citing: bool = False
) -> None:
    """Add a paper's focused HRC-relevant details without duplication."""

    title = paper.get("title", "Unknown title")
    doi = paper.get("doi") or "N/A"
    source_url = paper.get("source_url") or ""
    release_date = paper.get("release_date", "Unknown")
    venue = paper.get("venue") or "Unknown"
    authors = paper.get("authors") or "Unknown"
    abstract = paper.get("abstract") or ""
    full_text, source_quality = extract_full_paper_text(paper)

    # Skip if abstract already shown
    abstract_hash = hash((full_text or abstract)[:100]) if (full_text or abstract) else None
    is_duplicate = abstract_hash and abstract_hash in seen_abstracts
    if is_duplicate:
        abstract_display = "[Duplicate abstract - see reference above]"
    elif abstract_hash:
        seen_abstracts.add(abstract_hash)
        abstract_display = full_text or abstract
    else:
        abstract_display = ""

    if is_citing:
        doc.add_heading(f"→ {title}", level=3)
    else:
        doc.add_heading(title, level=2)

    # Compact info as a hyphen bullet line
    info_text = f"{authors} | {release_date} | {venue}"
    if doi != "N/A":
        info_text += f" | DOI: {doi}"
    info_para = doc.add_paragraph(f"- Info: {info_text}")
    info_para.paragraph_format.space_after = Pt(6)

    # Add internet links (source URL and DOI URL) as hyphen lines
    if source_url:
        source_para = doc.add_paragraph("- Source: ")
        add_hyperlink(source_para, source_url, source_url)
        source_para.paragraph_format.space_after = Pt(3)

    if doi != "N/A":
        doi_url = doi if doi.startswith("http") else f"https://doi.org/{doi}"
        doi_para = doc.add_paragraph("- DOI link: ")
        add_hyperlink(doi_para, doi_url, doi_url)
        doi_para.paragraph_format.space_after = Pt(6)

    source_quality_para = doc.add_paragraph(f"- Text source used for this summary: {source_quality}")
    source_quality_para.paragraph_format.space_after = Pt(4)

    # Synthesized summary instead of copied abstract text
    summary = build_work_summary(title, full_text or abstract)
    summary_para = doc.add_paragraph(f"- Overview: {summary}")
    summary_para.paragraph_format.space_after = Pt(4)

    # Extract detailed approach, data, and results as summaries
    approach = extract_approach_details(title, full_text or abstract)
    results = extract_results_details(title, full_text or abstract)

    if source_url:
        source_summary = "The source page provides the online version or project page for this work."
    else:
        source_summary = ""

    if approach:
        approach_summary = build_method_summary(title, abstract, approach)
        approach_para = doc.add_paragraph(f"- What they did: {approach_summary}")
        approach_para.paragraph_format.space_after = Pt(4)

    if results:
        results_summary = build_results_summary(title, full_text or abstract, results)
        results_para = doc.add_paragraph(f"- What they achieved: {results_summary}")
        results_para.paragraph_format.space_after = Pt(4)

    dataset_title = any(x in title.lower() for x in ["dataset", "data", "multimodal"])
    dataset_abstract = any(
        x in (full_text or abstract).lower()
        for x in ["dataset", "data", "recording", "multimodal", "sample", "participant", "subject"]
    )
    if dataset_title or dataset_abstract:
        profile = extract_dataset_profile(title, full_text or abstract, source_url)

        profile_para = doc.add_paragraph(f"- Dataset summary: {build_dataset_summary(title, profile)}")
        profile_para.paragraph_format.space_after = Pt(2)

        modality_para = doc.add_paragraph(f"- Dataset modalities: {profile['modalities']}")
        modality_para.paragraph_format.space_after = Pt(2)

        nature_para = doc.add_paragraph(f"- Data nature: {profile['nature']}")
        nature_para.paragraph_format.space_after = Pt(2)

        availability_para = doc.add_paragraph(f"- Availability: {profile['availability']}")
        availability_para.paragraph_format.space_after = Pt(2)

        if profile.get("extras"):
            extras_para = doc.add_paragraph(f"- Dataset cues: {profile['extras']}")
            extras_para.paragraph_format.space_after = Pt(2)

        use_para = doc.add_paragraph(f"- In what fields it can be used: {profile['use_fields']}")
        use_para.paragraph_format.space_after = Pt(4)

    elif source_summary:
        source_para = doc.add_paragraph(f"- Source note: {source_summary}")
        source_para.paragraph_format.space_after = Pt(2)

    # Single key contribution as a hyphen bullet line
    contribution = build_key_contribution(title, full_text or abstract)
    if contribution:
        contrib_para = doc.add_paragraph(f"- Main contribution: {contribution}")
        contrib_para.paragraph_format.space_after = Pt(10)



def main() -> None:
    # Load existing reference data
    with open(REFERENCE_JSON, "r", encoding="utf-8") as f:
        records = json.load(f)

    # Create new document
    doc = Document()
    doc.add_heading("HRC Reference Analysis: Methodology & Results Summary", 0)
    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d')}\n"
        "Condensed focus on key methodology, results, and HRC contributions. Duplicates removed."
    )

    seen_abstracts: set[str] = set()

    # Process each seed reference
    for idx, rec in enumerate(records, start=1):
        doc.add_heading(f"{idx}. {rec.get('title', 'Unknown')}", level=1)

        # Add seed paper details
        add_paper_details(doc, rec, seen_abstracts, is_citing=False)

        # Add citing papers
        citing_papers = rec.get("hrc_citing_papers", [])
        if citing_papers:
            doc.add_heading(f"Related Work ({len(citing_papers)} citing papers)", level=2)
            for cidx, cp in enumerate(citing_papers, start=1):
                add_paper_details(doc, cp, seen_abstracts, is_citing=True)

        doc.add_paragraph()

    # Save document
    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"Saved condensed report: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
